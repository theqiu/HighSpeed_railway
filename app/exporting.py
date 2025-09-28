"""Utilities for exporting dashboard artefacts in multiple formats."""

from __future__ import annotations

import re
from io import BytesIO
from typing import Dict, Mapping

import pandas as pd
import plotly.graph_objects as go
import plotly.io as pio
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas


def figure_to_image_bytes(fig: go.Figure, fmt: str = "jpg", scale: int = 2) -> bytes:
    """Render a Plotly figure to image bytes using Kaleido."""

    return pio.to_image(fig, format=fmt, scale=scale)


def _safe_sheet_name(name: str) -> str:
    sanitized = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_-]+", "_", name)[:31]
    return sanitized or "Sheet1"


def export_tables_to_excel(tables: Mapping[str, pd.DataFrame]) -> bytes:
    """Export multiple tables into a single Excel workbook."""

    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for title, table in tables.items():
            safe_name = _safe_sheet_name(title)
            table.to_excel(writer, sheet_name=safe_name, index=False)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_word(
    summary: Mapping[str, object],
    tables: Mapping[str, pd.DataFrame],
    images: Mapping[str, bytes],
) -> bytes:
    """Generate a Word report containing key metrics, tables, and charts."""

    document = Document()
    document.add_heading("检修数据分析报告", 0)

    if summary:
        document.add_heading("概要指标", level=1)
        summary_table = document.add_table(rows=len(summary) + 1, cols=2)
        summary_table.style = "Light Grid"
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = "指标"
        header_cells[1].text = "数值"
        for idx, (key, value) in enumerate(summary.items(), start=1):
            summary_table.rows[idx].cells[0].text = str(key)
            summary_table.rows[idx].cells[1].text = f"{value}"

    for title, table in tables.items():
        document.add_heading(title, level=1)
        preview = table.head(15)
        if preview.empty:
            document.add_paragraph("无数据")
            continue
        table_obj = document.add_table(rows=len(preview) + 1, cols=len(preview.columns))
        table_obj.style = "Light List"
        header = table_obj.rows[0].cells
        for col_idx, column in enumerate(preview.columns):
            header[col_idx].text = str(column)
        for row_idx, (_, row) in enumerate(preview.iterrows(), start=1):
            for col_idx, value in enumerate(row):
                table_obj.rows[row_idx].cells[col_idx].text = "" if pd.isna(value) else str(value)

    if images:
        document.add_heading("图表", level=1)
        for title, image_bytes in images.items():
            document.add_heading(title, level=2)
            document.add_picture(BytesIO(image_bytes), width=Inches(6))

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


def export_to_pdf(
    summary: Mapping[str, object],
    tables: Mapping[str, pd.DataFrame],
    images: Mapping[str, bytes],
) -> bytes:
    """Generate a PDF report containing metrics, sample tables, and charts."""

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 40

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(40, y, "检修数据分析报告")
    y -= 30

    pdf.setFont("Helvetica", 11)
    for key, value in summary.items():
        line = f"{key}: {value}"
        pdf.drawString(40, y, line[:120])
        y -= 18
        if y < 100:
            pdf.showPage()
            y = height - 40
            pdf.setFont("Helvetica", 11)

    pdf.setFont("Helvetica", 11)
    for title, table in tables.items():
        preview = table.head(10)
        if y < 120:
            pdf.showPage()
            y = height - 40
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(40, y, title)
        y -= 20
        pdf.setFont("Helvetica", 9)
        if preview.empty:
            pdf.drawString(40, y, "无数据")
            y -= 18
        else:
            for line in preview.to_string(index=False).splitlines():
                pdf.drawString(40, y, line[:150])
                y -= 14
                if y < 80:
                    pdf.showPage()
                    y = height - 40
                    pdf.setFont("Helvetica", 9)
            y -= 10

    for title, image_bytes in images.items():
        if y < 220:
            pdf.showPage()
            y = height - 40
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(40, y, title)
        y -= 20
        pdf.setFont("Helvetica", 9)
        image = ImageReader(BytesIO(image_bytes))
        img_width, img_height = image.getSize()
        scale = min((width - 80) / img_width, 300 / img_height)
        display_width = img_width * scale
        display_height = img_height * scale
        pdf.drawImage(image, 40, y - display_height, width=display_width, height=display_height)
        y -= display_height + 30

    pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


def prepare_chart_images(figures: Mapping[str, go.Figure], fmt: str = "png") -> Dict[str, bytes]:
    """Convert Plotly figures to image bytes for document exports."""

    images: Dict[str, bytes] = {}
    for title, fig in figures.items():
        try:
            images[title] = figure_to_image_bytes(fig, fmt=fmt)
        except Exception:
            # Skip figures that cannot be rendered in the current environment.
            continue
    return images
